#!/usr/bin/env python3
"""
Korean Markdown to Word (.docx) Converter with Embedded Images

한국어 번역된 학술 논문 Markdown을 이미지 포함 Word 문서로 변환합니다.

Usage:
    # 단일 파일
    python md_to_docx_kr.py --input 논문_kr.md --output 논문_kr.docx --image-base ./md_converted/

    # 배치 모드
    python md_to_docx_kr.py --input-dir ./translated/ --output-dir ./translated_output/ --image-base ./md_converted/

Features:
    - 한국어 폰트 (맑은 고딕) + 영문 폰트 (Times New Roman)
    - 이미지 자동 삽입 (페이지 폭에 맞춤)
    - Markdown 표 → Word 표 변환
    - Figure/Table 캡션 중앙 정렬
    - Vancouver 참고문헌 서식

Requirements:
    pip install python-docx Pillow
"""

import argparse
import os
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# 한국어 폰트 설정
KR_FONT = "맑은 고딕"
EN_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(16)
HEADING2_SIZE = Pt(14)
HEADING3_SIZE = Pt(12)
CAPTION_SIZE = Pt(10)
REF_SIZE = Pt(9)
LINE_SPACING = 1.5
MAX_IMAGE_WIDTH_CM = 16.0


def create_document() -> Document:
    """한국어 학술 문서 스타일 생성"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = EN_FONT
    font.size = BODY_SIZE
    # 한국어 폰트 설정 (rFonts)
    rpr = style.element.rPr
    if rpr is None:
        rpr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = style.element.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), KR_FONT)

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    return doc


def set_run_font(run, bold=False, italic=False, size=None, color=None):
    """Run에 한국어+영문 폰트 동시 설정"""
    run.font.name = EN_FONT
    run.font.size = size or BODY_SIZE
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # eastAsia 폰트
    r = run._element
    rpr = r.find(qn('w:rPr'))
    if rpr is None:
        rpr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = r.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), KR_FONT)


def parse_inline(paragraph, text: str, size=None):
    """인라인 서식 파싱 (굵은 글씨, 이탤릭, 참고문헌 번호)"""
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'        # bold
        r'|(\*(.+?)\*)'           # italic
        r'|(\[[\d,\s\-–]+\])'    # citation numbers
        r'|([^*\[]+)'             # plain text
        r'|(.)'                   # fallback
    )

    for m in pattern.finditer(text):
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            set_run_font(run, bold=True, size=size)
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            set_run_font(run, italic=True, size=size)
        elif m.group(5):  # citation
            run = paragraph.add_run(m.group(5))
            set_run_font(run, size=size)
        elif m.group(6):  # plain
            run = paragraph.add_run(m.group(6))
            set_run_font(run, size=size)
        elif m.group(7):  # fallback
            run = paragraph.add_run(m.group(7))
            set_run_font(run, size=size)


def find_image_file(image_ref: str, image_base: Path, md_stem: str) -> Optional[Path]:
    """이미지 파일 경로를 다양한 위치에서 탐색

    탐색 순서:
    1. image_base / image_ref (직접 경로)
    2. image_base / md_stem / images / filename
    3. image_base / images / filename (구 경로)
    4. 원본 stem으로도 시도 (_kr 제거)
    """
    # angle bracket 제거
    clean_ref = image_ref.strip('<>')

    # 1. 직접 경로
    direct = image_base / clean_ref
    if direct.exists():
        return direct

    # 2. stem 기반 (번역 파일명에서 _kr 제거)
    orig_stem = md_stem
    if orig_stem.endswith('_kr'):
        orig_stem = orig_stem[:-3]

    filename = Path(clean_ref).name

    for stem in [orig_stem, md_stem]:
        candidate = image_base / stem / "images" / filename
        if candidate.exists():
            return candidate

    # 3. images/ 직접
    candidate = image_base / "images" / filename
    if candidate.exists():
        return candidate

    # 4. 하위 폴더 탐색
    for img_dir in image_base.rglob("images"):
        candidate = img_dir / filename
        if candidate.exists():
            return candidate

    return None


def add_image_to_doc(doc: Document, image_path: Path):
    """이미지를 문서에 삽입 (페이지 폭에 맞춤, doc.add_picture 사용)"""
    try:
        from PIL import Image
        with Image.open(str(image_path)) as img:
            w_px, h_px = img.size

        if w_px == 0 or h_px == 0:
            return

        # 이미지 DPI 기본값 96
        w_inches = w_px / 96.0
        h_inches = h_px / 96.0

        max_w_inches = MAX_IMAGE_WIDTH_CM / 2.54

        # 최대 폭에 맞춤
        if w_inches > max_w_inches:
            ratio = max_w_inches / w_inches
            w_inches *= ratio
            h_inches *= ratio

        # 극단적 비율 이미지(수식 등): 최소 높이 0.3인치 보장
        if h_inches < 0.3:
            h_inches = 0.3

        # doc.add_picture() 사용 (run.add_picture보다 Word 호환성 우수)
        doc.add_picture(str(image_path),
                        width=Inches(w_inches), height=Inches(h_inches))
        # add_picture가 추가한 마지막 paragraph에 중앙 정렬 적용
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    except ImportError:
        # Pillow 없으면 기본 크기로 삽입
        doc.add_picture(str(image_path), width=Cm(MAX_IMAGE_WIDTH_CM))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        # 이미지 삽입 실패 시 경고 텍스트
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[이미지 삽입 실패: {image_path.name} — {e}]")
        set_run_font(run, italic=True, size=CAPTION_SIZE,
                     color=RGBColor(0x99, 0x99, 0x99))


def is_abstract_table(table_lines: list) -> bool:
    """초록/논문정보 표인지 감지

    pymupdf4llm이 생성한 초록 표 패턴:
    |논문 정보<br>키워드:<br>...|초 록||---|---|
    ||본문 텍스트...|
    """
    if not table_lines:
        return False
    first_line = table_lines[0].lower()
    has_info = any(kw in first_line for kw in [
        '논문 정보', '논 문 정 보', '논문정보', 'article info',
    ])
    has_abstract = any(kw in first_line for kw in [
        '초 록', '초록', 'abstract',
    ])
    return has_info and has_abstract


def parse_abstract_table(table_lines: list) -> dict:
    """초록 표에서 논문 정보 라벨, 키워드, 초록 본문 추출"""
    result = {'info_label': '논문 정보', 'keywords': [], 'abstract_label': '초록',
              'abstract': ''}

    rows = parse_markdown_table(table_lines)
    if not rows:
        return result

    # 첫 행: 논문 정보 + 키워드 셀, 초록 헤더
    first_row = rows[0]
    if first_row:
        info_cell = first_row[0]
        parts = re.split(r'<br\s*/?>', info_cell)
        if parts:
            result['info_label'] = parts[0].strip()
        for part in parts[1:]:
            part = part.strip()
            if not part:
                continue
            # "키워드:" 또는 "Keywords:" 라벨은 건너뛰기
            if re.match(r'^(키워드|keywords?)\s*:?\s*$', part, re.IGNORECASE):
                continue
            result['keywords'].append(part)

        # 초록 헤더 라벨 찾기
        for cell in first_row[1:]:
            cell = cell.strip()
            if cell and '---' not in cell:
                result['abstract_label'] = cell
                break

    # 이후 행에서 초록 본문 추출 (가장 긴 텍스트)
    for row in rows[1:]:
        for cell in row:
            cell = cell.strip()
            if cell and len(cell) > len(result['abstract']):
                result['abstract'] = cell

    return result


def add_abstract_to_doc(doc: Document, info: dict, result: dict):
    """초록/논문정보를 본문 형식으로 문서에 추가"""
    # 키워드
    if info['keywords']:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run('키워드: ')
        set_run_font(run, bold=True, size=BODY_SIZE)
        run = p.add_run(', '.join(info['keywords']))
        set_run_font(run, size=BODY_SIZE)
        result["paragraphs"] += 1

    # 초록 제목
    if info['abstract']:
        heading = doc.add_heading(info['abstract_label'], level=2)
        for run in heading.runs:
            set_run_font(run, bold=True, size=HEADING2_SIZE)
        result["headings"] += 1

        # 초록 본문
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        parse_inline(p, info['abstract'])
        result["paragraphs"] += 1


def parse_markdown_table(lines: list) -> list:
    """Markdown 표를 2D 리스트로 파싱"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|') or not line.endswith('|'):
            continue
        # 구분선 건너뛰기
        if re.match(r'^\|[\s\-:]+\|', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def add_table_to_doc(doc: Document, rows: list):
    """Word 표 추가 (Word 최대 63열 제한 준수)"""
    if not rows:
        return

    n_cols = max(len(r) for r in rows)

    # Word 최대 열 수는 63. 초과 시 깨진 표이므로 일반 텍스트로 삽입
    MAX_WORD_COLS = 63
    if n_cols > MAX_WORD_COLS:
        for row_data in rows:
            text = " | ".join(c for c in row_data if c)
            if text.strip():
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_run_font(run, size=Pt(9))
        doc.add_paragraph()
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text)
                set_run_font(run, bold=(i == 0), size=Pt(10))

    # 표 앞뒤 빈 줄
    doc.add_paragraph()


def convert_md_to_docx_kr(
    input_path: Path,
    output_path: Path,
    image_base: Optional[Path] = None
) -> dict:
    """한국어 Markdown → Word 변환"""
    content = input_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    md_stem = input_path.stem

    if image_base is None:
        image_base = input_path.parent

    doc = create_document()
    result = {
        "input": str(input_path),
        "output": str(output_path),
        "status": "pending",
        "paragraphs": 0,
        "headings": 0,
        "images_inserted": 0,
        "images_missing": 0,
        "tables": 0,
        "references": 0,
        "warnings": []
    }

    in_references = False
    in_table = False
    table_lines = []
    paragraph_buffer = []
    in_code_block = False

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = ' '.join(paragraph_buffer).strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.75)
                parse_inline(p, text)
                result["paragraphs"] += 1
            paragraph_buffer = []

    def flush_table():
        nonlocal table_lines, in_table
        if table_lines:
            if is_abstract_table(table_lines):
                info = parse_abstract_table(table_lines)
                add_abstract_to_doc(doc, info, result)
            else:
                rows = parse_markdown_table(table_lines)
                if rows:
                    add_table_to_doc(doc, rows)
                    result["tables"] += 1
            table_lines = []
        in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 코드 블록
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # 빈 줄
        if not stripped:
            if in_table:
                flush_table()
            flush_paragraph()
            i += 1
            continue

        # HTML 주석 건너뛰기
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            i += 1
            continue

        # 표 처리
        if stripped.startswith('|') and stripped.endswith('|'):
            flush_paragraph()
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # 이미지 참조
        img_match = re.match(r'^!\[([^\]]*)\]\((.+?)\)\s*$', stripped)
        if img_match:
            flush_paragraph()
            alt_text = img_match.group(1)
            img_ref = img_match.group(2)

            img_path = find_image_file(img_ref, image_base, md_stem)
            if img_path:
                add_image_to_doc(doc, img_path)
                result["images_inserted"] += 1
            else:
                result["images_missing"] += 1
                result["warnings"].append(f"Image not found: {img_ref}")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[이미지 누락: {img_ref}]")
                set_run_font(run, italic=True, size=CAPTION_SIZE,
                             color=RGBColor(0x99, 0x99, 0x99))
            i += 1
            continue

        # 인라인 이미지 (텍스트 중간에 ![...](...)가 있는 경우)
        if '![' in stripped and not stripped.startswith('!'):
            # 이미지 전후 텍스트 분리
            parts = re.split(r'(!\[[^\]]*\]\([^)]+\))', stripped)
            flush_paragraph()
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                img_inline = re.match(r'^!\[([^\]]*)\]\((.+?)\)$', part)
                if img_inline:
                    img_ref = img_inline.group(2)
                    img_path = find_image_file(img_ref, image_base, md_stem)
                    if img_path:
                        add_image_to_doc(doc, img_path)
                        result["images_inserted"] += 1
                    else:
                        result["images_missing"] += 1
                else:
                    p = doc.add_paragraph()
                    parse_inline(p, part)
                    result["paragraphs"] += 1
            i += 1
            continue

        # 헤더
        header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if header_match:
            flush_paragraph()
            level = len(header_match.group(1))
            header_text = header_match.group(2).strip()

            if any(kw in header_text.lower() for kw in ['references', 'reference', '참고문헌']):
                in_references = True

            heading_level = max(1, level - 1)
            heading = doc.add_heading(header_text, level=heading_level)
            for run in heading.runs:
                set_run_font(run, bold=True,
                             size=[HEADING1_SIZE, HEADING2_SIZE, HEADING3_SIZE, HEADING3_SIZE][min(level-1, 3)])
            result["headings"] += 1
            i += 1
            continue

        # Figure/Table 캡션 (중앙 정렬)
        caption_match = re.match(
            r'^\*\*(그림|표|Figure|Table)\s*\d+[\.\):]?\*\*', stripped)
        if caption_match:
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_inline(p, stripped, size=CAPTION_SIZE)
            result["paragraphs"] += 1
            i += 1
            continue

        # References
        if in_references:
            flush_paragraph()
            ref_match = re.match(r'^\[?(\d+)\]?\s*(.+)$', stripped)
            if ref_match:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(1.0)
                text = f'[{ref_match.group(1)}] {ref_match.group(2)}'
                run = p.add_run(text)
                set_run_font(run, size=REF_SIZE)
                result["references"] += 1
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(stripped)
                set_run_font(run, size=REF_SIZE)
            i += 1
            continue

        # 일반 텍스트
        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()
    if in_table:
        flush_table()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    result["status"] = "success"

    return result


def batch_convert(
    input_dir: Path,
    output_dir: Path,
    image_base: Optional[Path] = None
) -> dict:
    """배치 변환"""
    output_dir.mkdir(parents=True, exist_ok=True)
    md_files = sorted(input_dir.glob("*_kr.md"))

    if not md_files:
        md_files = sorted(input_dir.glob("*.md"))

    print(f"Found {len(md_files)} files to convert")
    print("-" * 50)

    results = []
    success = 0

    for i, md_path in enumerate(md_files, 1):
        print(f"[{i}/{len(md_files)}] {md_path.name}")
        out_path = output_dir / md_path.with_suffix('.docx').name

        r = convert_md_to_docx_kr(md_path, out_path, image_base)
        results.append(r)

        if r["status"] == "success":
            success += 1
            size_kb = out_path.stat().st_size / 1024
            print(f"  -> OK ({size_kb:.0f} KB, {r['images_inserted']} images, "
                  f"{r['tables']} tables)")
            if r["warnings"]:
                for w in r["warnings"][:3]:
                    print(f"  -> WARNING: {w}")
        else:
            print(f"  -> FAILED: {r.get('error', 'Unknown')}")

    print("-" * 50)
    print(f"Complete: {success}/{len(md_files)} converted")

    return {
        "summary": {"total": len(md_files), "success": success,
                     "failed": len(md_files) - success},
        "files": results
    }


def main():
    parser = argparse.ArgumentParser(
        description="Convert Korean translated Markdown to Word with images"
    )
    parser.add_argument("--input", "-i", type=Path, default=None,
                        help="Single input MD file")
    parser.add_argument("--output", "-o", type=Path, default=None,
                        help="Output .docx path")
    parser.add_argument("--input-dir", type=Path, default=None,
                        help="Batch: input directory with *_kr.md files")
    parser.add_argument("--output-dir", type=Path, default=None,
                        help="Batch: output directory")
    parser.add_argument("--image-base", type=Path, default=None,
                        help="Base directory for image lookup")

    args = parser.parse_args()

    if args.input:
        if not args.input.exists():
            print(f"Error: {args.input} not found")
            return 1
        output = args.output or args.input.with_suffix('.docx')
        r = convert_md_to_docx_kr(args.input, output, args.image_base)
        print(f"Status: {r['status']}")
        print(f"Paragraphs: {r['paragraphs']}, Headings: {r['headings']}")
        print(f"Images: {r['images_inserted']} inserted, {r['images_missing']} missing")
        print(f"Tables: {r['tables']}, References: {r['references']}")
        if r["warnings"]:
            print("Warnings:")
            for w in r["warnings"]:
                print(f"  - {w}")
    elif args.input_dir:
        if not args.input_dir.exists():
            print(f"Error: {args.input_dir} not found")
            return 1
        out_dir = args.output_dir or args.input_dir / "docx_output"
        batch_convert(args.input_dir, out_dir, args.image_base)
    else:
        parser.error("Either --input or --input-dir is required")

    return 0


if __name__ == "__main__":
    exit(main())

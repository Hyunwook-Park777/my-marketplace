#!/usr/bin/env python3
"""
ISO DIS Ballot Form Writer

빈 ISO DIS 투표양식(Word)에 검토 의견을 자동 삽입합니다.
ballot_content.json의 내용을 Word 템플릿의 적절한 위치에 삽입합니다.

Usage:
    python ballot_writer.py --template "ISO-DIS_8178-1.docx" --content ballot_content.json -o ballot_filled.docx

ballot_content.json 구조:
{
    "standard_name": "ISO/DIS 8178-1",
    "reviewer": { "name": "...", "organization": "...", "date": "..." },
    "vote": "approval_with_comments",
    "review_summary": "...",
    "key_changes": [ { "clause": "...", "description": "...", "opinion": "..." } ],
    "technical_comments": [ { "clause": "...", "current_text": "...", "proposed_text": "...", "reason": "..." } ],
    "general_comments": "...",
    "recommendation": "..."
}
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("에러: python-docx가 필요합니다. pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 투표 선택 매핑
VOTE_MAP = {
    "approval": "Approval (승인)",
    "approval_with_comments": "Approval with comments (조건부 승인)",
    "disapproval": "Disapproval (반대)",
    "abstention": "Abstention (기권)",
}

# 투표 키워드 → 문서 내 매칭 텍스트
VOTE_MARKERS = {
    "approval": ["Approval", "승인", "찬성"],
    "approval_with_comments": ["Approval with comments", "조건부 승인", "조건부 찬성"],
    "disapproval": ["Disapproval", "반대"],
    "abstention": ["Abstention", "기권"],
}


def find_paragraph_index(doc, search_texts, start_from=0):
    """문서에서 검색 텍스트를 포함하는 단락의 인덱스를 반환"""
    for i, para in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        text = para.text.strip()
        for search in search_texts:
            if search.lower() in text.lower():
                return i
    return -1


def set_paragraph_font(paragraph, font_name="맑은 고딕", font_size=11):
    """단락의 기본 글꼴 설정"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)


def add_formatted_paragraph(doc, text, font_name="맑은 고딕", font_size=11,
                           bold=False, alignment=None, space_after=None,
                           insert_before_idx=None):
    """서식이 적용된 단락을 추가하거나 지정 위치에 삽입"""
    if insert_before_idx is not None and insert_before_idx < len(doc.paragraphs):
        # 지정 위치 앞에 삽입
        ref_para = doc.paragraphs[insert_before_idx]
        new_para = ref_para.insert_paragraph_before(text)
    else:
        new_para = doc.add_paragraph(text)

    run = new_para.runs[0] if new_para.runs else new_para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

    if alignment is not None:
        new_para.alignment = alignment

    if space_after is not None:
        new_para.paragraph_format.space_after = Pt(space_after)

    return new_para


def build_review_content(content: dict) -> str:
    """ballot_content.json으로부터 문서검토내용 텍스트를 생성"""
    lines = []

    # 검토자 정보
    reviewer = content.get("reviewer", {})
    if reviewer:
        lines.append(f"검토자: {reviewer.get('name', '')}")
        if reviewer.get('organization'):
            lines.append(f"소속: {reviewer['organization']}")
        lines.append(f"검토일: {reviewer.get('date', datetime.now().strftime('%Y-%m-%d'))}")
        lines.append("")

    # 표준 정보
    lines.append(f"대상 표준: {content.get('standard_name', '')}")
    lines.append(f"투표 권고: {VOTE_MAP.get(content.get('vote', ''), content.get('vote', ''))}")
    lines.append("")

    # 검토 요약
    if content.get("review_summary"):
        lines.append("=" * 60)
        lines.append("1. 개정 개요 및 검토 의견 요약")
        lines.append("=" * 60)
        lines.append("")
        lines.append(content["review_summary"])
        lines.append("")

    # 핵심 변경 사항
    key_changes = content.get("key_changes", [])
    if key_changes:
        lines.append("=" * 60)
        lines.append("2. 핵심 변경사항 분석")
        lines.append("=" * 60)
        lines.append("")
        for i, change in enumerate(key_changes, 1):
            lines.append(f"  {i}) 조항 {change.get('clause', 'N/A')}")
            lines.append(f"     변경 내용: {change.get('description', '')}")
            lines.append(f"     검토 의견: {change.get('opinion', '')}")
            lines.append("")

    # 기술적 의견 (comment 형식)
    tech_comments = content.get("technical_comments", [])
    if tech_comments:
        lines.append("=" * 60)
        lines.append("3. 기술적 의견 (Technical Comments)")
        lines.append("=" * 60)
        lines.append("")
        for i, comment in enumerate(tech_comments, 1):
            lines.append(f"  Comment {i}:")
            lines.append(f"    Clause: {comment.get('clause', 'N/A')}")
            if comment.get('current_text'):
                lines.append(f"    현재 문구: {comment['current_text']}")
            if comment.get('proposed_text'):
                lines.append(f"    제안 변경: {comment['proposed_text']}")
            lines.append(f"    근거: {comment.get('reason', '')}")
            lines.append("")

    # 일반 의견
    if content.get("general_comments"):
        lines.append("=" * 60)
        lines.append("4. 일반 의견 (General Comments)")
        lines.append("=" * 60)
        lines.append("")
        lines.append(content["general_comments"])
        lines.append("")

    # 최종 권고
    if content.get("recommendation"):
        lines.append("=" * 60)
        lines.append("5. 최종 권고사항")
        lines.append("=" * 60)
        lines.append("")
        lines.append(content["recommendation"])
        lines.append("")

    return '\n'.join(lines)


def fill_ballot_form(template_path: str, content: dict, output_path: str):
    """투표양식 Word 파일에 검토 의견을 삽입"""
    doc = Document(template_path)

    # 1. 문서검토내용 섹션 찾기
    review_section_idx = find_paragraph_index(
        doc,
        ["문서검토내용", "검토 내용", "Review content", "Reviewer's comments"]
    )

    # 2. 투표 섹션 찾기
    vote_section_idx = find_paragraph_index(
        doc,
        ["Do you approve", "approve the document", "투표", "승인 여부"]
    )

    # 검토 내용 생성
    review_text = build_review_content(content)

    # 3. 검토 내용 삽입
    if review_section_idx >= 0:
        # 검토 섹션 바로 다음에 삽입
        insert_idx = review_section_idx + 1

        # 투표 섹션 전까지의 빈 단락들 사이에 삽입
        for line in review_text.split('\n'):
            if insert_idx < len(doc.paragraphs):
                ref_para = doc.paragraphs[insert_idx]
                new_para = ref_para.insert_paragraph_before(line)
            else:
                new_para = doc.add_paragraph(line)

            # 서식 적용
            if new_para.runs:
                run = new_para.runs[0]
            else:
                run = new_para.add_run("")
            run.font.name = "맑은 고딕"
            run.font.size = Pt(11)

            # 섹션 제목은 볼드
            if line.startswith("=") or (line and line[0].isdigit() and '.' in line[:3]):
                run.bold = True
    else:
        # 검토 섹션을 찾지 못한 경우, 문서 끝에 추가
        doc.add_paragraph("")
        heading_para = doc.add_paragraph("문서검토내용 및 검토자 의견")
        if heading_para.runs:
            heading_para.runs[0].bold = True
            heading_para.runs[0].font.size = Pt(14)

        for line in review_text.split('\n'):
            para = doc.add_paragraph(line)
            if para.runs:
                para.runs[0].font.name = "맑은 고딕"
                para.runs[0].font.size = Pt(11)

    # 4. 투표 선택 표시
    vote = content.get("vote", "approval_with_comments")
    if vote_section_idx >= 0:
        # 투표 섹션 이후에서 해당 투표 옵션 찾아 표시
        markers = VOTE_MARKERS.get(vote, [])
        for i, para in enumerate(doc.paragraphs):
            if i <= vote_section_idx:
                continue
            text = para.text.strip()
            for marker in markers:
                if marker.lower() in text.lower():
                    # 체크 표시 추가 (☑ 또는 ■)
                    if "☐" in para.text:
                        for run in para.runs:
                            run.text = run.text.replace("☐", "☑")
                    elif "□" in para.text:
                        for run in para.runs:
                            run.text = run.text.replace("□", "■")
                    else:
                        # 앞에 체크 표시 추가
                        if para.runs:
                            para.runs[0].text = "☑ " + para.runs[0].text
                    break

    # 저장
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser(
        description='ISO DIS 투표양식 자동 작성'
    )
    parser.add_argument('--template', '-t', required=True,
                        help='투표양식 Word 템플릿 경로')
    parser.add_argument('--content', '-c', required=True,
                        help='검토 내용 JSON 파일 경로 (ballot_content.json)')
    parser.add_argument('--output', '-o', required=True,
                        help='출력 Word 파일 경로')

    args = parser.parse_args()

    try:
        # 템플릿 확인
        template_path = Path(args.template)
        if not template_path.exists():
            raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {args.template}")

        # 내용 로드
        content_path = Path(args.content)
        if not content_path.exists():
            raise FileNotFoundError(f"내용 파일을 찾을 수 없습니다: {args.content}")

        with open(content_path, 'r', encoding='utf-8') as f:
            content = json.load(f)

        # 투표양식 작성
        fill_ballot_form(str(template_path), content, args.output)
        print(f"투표양식 작성 완료: {args.output}")

        # 요약 출력
        vote = VOTE_MAP.get(content.get("vote", ""), content.get("vote", ""))
        n_tech = len(content.get("technical_comments", []))
        n_changes = len(content.get("key_changes", []))
        print(f"  투표: {vote}")
        print(f"  핵심 변경 분석: {n_changes}건")
        print(f"  기술적 의견: {n_tech}건")

    except Exception as e:
        print(f"에러: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
